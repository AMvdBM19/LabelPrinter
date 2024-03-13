from docx import Document

from docx.shared import Pt

from docx.enum.section import WD_ORIENT

from docx.shared import Inches

from docx.shared import Mm

import time

import easygui

import os

from colorama import Fore, Back, Style

from colorama import just_fix_windows_console

from docx.enum.text import WD_ALIGN_PARAGRAPH

just_fix_windows_console()

document = Document()

section = document.sections


print("")

print(Fore.GREEN + "Welcome to GPC Universal Label Creator")

time.sleep(1)

print(Fore.WHITE + "")


#Font type parameters
style = document.styles['Normal']
font = style.font
font.name = 'Arial'


print("What format of labels do you wish to create?")
print("")
print("Please choose between UPS size labels (4inch x 6inch), 1.4inch x 1.6inch labels and 1.18inch x 0.79inch labels.")
print("")
time.sleep(1)

#label_format input parameters

small = ("small", "s", "klein", "k")
medium =("medium", "m", "middel")
large = ("large", "l", "groot", "g")


while True:
    
    print("Type ''large'' for 4inch x 6inch labels; ''medium'' for 1.4inch x 1.6inch labels and ''small''")
    print("for 1.18inch x 0.79inch or ''exit'' if you want to close the program.")
    print("")

    label_format = input(":").lower()
    print("")
    label_format = str(label_format)

    
    if label_format in large:
        
        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        
        section.page_height = Mm(101.6)
        section.page_width = Mm(152)

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(42.4)
        section.bottom_margin = Mm(25.4)

        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        font.size = Pt(72)
        
        break

    elif label_format in medium:

        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        
        section.page_height = Mm(35.0)
        section.page_width = Mm(37.5)

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(5.0)
        section.bottom_margin = Mm(5.0)

        section.header_distance = Mm(2.7)
        section.footer_distance = Mm(2.7)

        font.size = Pt(22)
        
        break

    elif label_format in small:

        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        
        section.page_height = Mm(17.5)
        section.page_width = Mm(27.5)

        section.left_margin = Mm(0)
        section.right_margin = Mm(0)
        section.top_margin = Mm(5.0)
        section.bottom_margin = Mm(2.0)

        section.header_distance = Mm(1.7)
        section.footer_distance = Mm(1.7)

        font.size = Pt(16)
        
        break

    elif label_format == "exit" or label_format == "quit":
        
        time.sleep(2)
        
        break

    else:
        
        print(Fore.RED + "User input not valid")
        print(Fore.WHITE + "")
        time.sleep(2)
        
        continue


while True:
    
    if label_format == "exit" or label_format == "quit":
        
        time.sleep(1)
        print(Fore.GREEN + "Thanks for using GPC Label Creator")
        print(Fore.WHITE + "")
        time.sleep(2)
        quit()

    else:
        
        time.sleep(1)
        print(Fore.WHITE + "")
        
        print("If you wish to print out a label for a particular SKU; insert the SKU letter code in the Bay input,")
        print("the numerical code in the column input and leave the amount of locations empty.")
        print("")

        letter = input("Please insert the bay: ").upper()
        
        if letter == "EXIT" or letter == "QUIT":
            time.sleep(1)
            print(Fore.GREEN + "Thanks for using GPC Label Creator")
            print(Fore.WHITE + "")
            time.sleep(1)
            quit()
        
        rack = input("Please insert the column within the bay: ")
        
        if rack == "exit" or rack == "quit":
            time.sleep(1)
            print(Fore.GREEN + "Thanks for using GPC Label Creator")
            print(Fore.WHITE + "")
            time.sleep(2)
            quit()

    try:

        amount_of_locations = input("Please insert the amount of locations: ")
        print("")
        amount_of_locations = int(amount_of_locations)
        
        if amount_of_locations < 0:

            print("")
            print(Fore.RED + "The amount of locations must be equal or greater than 0")
            print(Fore.WHITE + "")
            time.sleep(2)

            continue


        elif amount_of_locations > 201:

            print("")
            print(Fore.RED + "There is no capacity for more than 200 labels")
            print(Fore.WHITE + "")
            time.sleep(1)

            continue


        elif amount_of_locations == 0:

            label_list = (letter + rack + (" - ") + str(0))

            print(Fore.YELLOW + label_list)

            p = document.add_paragraph()

            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r = p.add_run(label_list)

            r.bold = True

            time.sleep(0.10)

            document.save("Label For " + str(letter) + str(rack) + ".docx")

            print(Fore.GREEN + "")

            
            

        print(Fore.WHITE + "")

        for result in (range (1, amount_of_locations + 1)):

            label_list = (letter + rack + (" - ") + str(result))

            print(Fore.YELLOW + label_list)

            p = document.add_paragraph()

            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            r = p.add_run(label_list)

            r.bold = True
        
            time.sleep(0.10)

            document.save("Labels For " + str(letter) + str(rack) + ".docx")

            print(Fore.GREEN + "")

            

    except:
        
        if  amount_of_locations == "":

            amount_of_locations == str
            label_list = (letter + rack)

            print(Fore.YELLOW + label_list)

            if label_format in large:
                
                font.size = Pt(72)
                section.top_margin = Mm(35.4)
                section.bottom_margin = Mm(25.4)
            
            elif label_format in medium:
                
                font.size = Pt(18)


            elif label_format in small:
                
                font.size = Pt(12)


            while True:

                print(Fore.WHITE)

                copies = input(Fore.WHITE + "How many copies do you wish to create?: ")
                print("")

                try:                               

                    if int(copies) > 200 :
                        time.sleep(1)
                        print(Fore.RED + "There is not enough capacity for more than 200 labels")
                        continue

                    elif int(copies) == 0:
                        time.sleep(1)
                        print(Fore.RED + "The amount of copies must be greater than 0")
                        continue


                    elif int(copies) < 0 :
                        time.sleep(1)
                        print(Fore.RED + "The amount of copies must be greater than 0")
                        continue
                
                    else:

                        for result1 in range (0, (int(copies))):

                            label_list1 = (letter + rack)

                            print(Fore.YELLOW + label_list1)

                            p = document.add_paragraph()

                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            r = p.add_run(label_list1)

                            r.bold = True
        
                            time.sleep(0.10)

                            document.save("Labels For " + label_list1 + ".docx")

                            print(Fore.GREEN + "")
                        break       
                
                except:
                    
                    copies = str(copies)

                    if str(copies) == "exit" or  str(copies) == "quit":

                        print("")
                        time.sleep(1)
                        print(Fore.GREEN + "Thanks for using GPC Label Creator")
                        time.sleep(1)
                        print(Fore.WHITE)
                        quit()
                        

                    else:
                        print()
                        print(Fore.RED + "Please select a valid amount of copies.")
                        print()
                        continue

        else:
        
            time.sleep(1)
            print("")
            print(Fore.RED + "Please select a valid amount to define the amount of locations")
            print(Fore.WHITE + "")
            time.sleep(1)

            continue    
    break


time.sleep(1)
print("The document Labels For " + "-" + str(letter) + str(rack) + "-" + " has been created succesfully")
print(Fore.WHITE + "")
print(Fore.GREEN + "Thanks for using GPC Label Creator")
print(Fore.WHITE + "")
time.sleep(2)
quit()

#By Andres van den Bos 2022
